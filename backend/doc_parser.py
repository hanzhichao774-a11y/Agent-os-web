"""公共文档解析器：读取 PDF、DOCX、XLSX、CSV、纯文本内容为字符串。"""

from pathlib import Path


def read_document_text(fpath: Path) -> str:
    """读取文档内容，支持 PDF、DOCX、XLSX、CSV、纯文本。"""
    suffix = fpath.suffix.lower()
    try:
        if suffix == ".pdf":
            import fitz
            pdf = fitz.open(str(fpath))
            parts = []
            for page in pdf:
                t = page.get_text().strip()
                if t:
                    parts.append(t)
            pdf.close()
            return "\n\n".join(parts)
        elif suffix == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(fpath))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)
        elif suffix in (".xlsx", ".xls"):
            from openpyxl import load_workbook
            wb = load_workbook(str(fpath), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"[工作表: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(parts)
        else:
            return fpath.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"[DOC_PARSER] 文件读取失败 {fpath.name}: {e}")
        return ""
