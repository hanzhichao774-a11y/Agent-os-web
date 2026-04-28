import io
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, UploadFile, File, Query

from config import KNOWLEDGE_DOCS_DIR
from knowledge import ingest_document, _uploaded_docs, list_documents
from agents import invalidate_agent
from database import register_task_file

router = APIRouter()


@router.post("/api/knowledge/upload")
async def api_upload_document(
    file: UploadFile = File(...),
    project_id: Optional[str] = Query(None),
    task_id: Optional[str] = Query(None),
):
    content = await file.read()
    doc_name = file.filename or "unknown.txt"
    suffix = Path(doc_name).suffix.lower()

    doc_path = KNOWLEDGE_DOCS_DIR / doc_name

    if suffix == ".pdf":
        import fitz
        doc_path.write_bytes(content)
        pdf = fitz.open(stream=content, filetype="pdf")
        _end_puncts = set('。；！？.;!?\n')
        parts: list[str] = []
        for page in pdf:
            t = page.get_text().strip()
            if not t:
                continue
            if parts and parts[-1] and parts[-1][-1] not in _end_puncts:
                parts.append(' ' + t)
            else:
                parts.append('\n\n' + t if parts else t)
        pdf.close()
        text = ''.join(parts).strip()

    elif suffix == ".docx":
        from docx import Document as DocxDocument
        doc_path.write_bytes(content)
        doc = DocxDocument(io.BytesIO(content))
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)

    elif suffix in (".xlsx", ".xls"):
        from openpyxl import load_workbook
        doc_path.write_bytes(content)
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"[工作表: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        text = "\n".join(parts)

    elif suffix == ".csv":
        text = content.decode("utf-8", errors="ignore")
        doc_path.write_text(text, encoding="utf-8")

    else:
        text = content.decode("utf-8", errors="ignore")
        doc_path.write_text(text, encoding="utf-8")

    if not text.strip():
        return {"success": False, "error": "文件内容为空，无法解析"}

    try:
        chunk_count = ingest_document(doc_name, text)
        _uploaded_docs[doc_name] = chunk_count
        invalidate_agent("a2")
        if project_id:
            register_task_file(project_id, task_id, doc_name, "upload", "knowledge")
        return {"success": True, "doc_name": doc_name, "chunks": chunk_count}
    except Exception as e:
        return {"success": False, "error": f"解析失败: {e}"}


@router.get("/api/knowledge/docs")
async def api_list_docs():
    return list_documents()
